from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("templates/pandoc_reference_tcc.docx")


def set_font(style, size_pt=12, bold=False, italic=False, color="000000"):
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(style, *, alignment=None, first_line=None, left=None, before=0, after=0, line_spacing=1.5):
    pf = style.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    if left is not None:
        pf.left_indent = Cm(left)


def add_or_get_style(document, name, kind):
    styles = document.styles
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, kind)


def set_cell_margins(style, margin_twips=120):
    tbl_pr = style._element.get_or_add_pPr()
    # Table style cell margins are not consistently honored by all consumers, but
    # this keeps Pandoc output from looking glued to borders in Word.
    tbl_style_pr = OxmlElement("w:tblPr")
    tbl_cell_mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")
        tbl_cell_mar.append(node)
    tbl_style_pr.append(tbl_cell_mar)
    tbl_pr.append(tbl_style_pr)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    styles = doc.styles

    normal = styles["Normal"]
    set_font(normal, 12)
    set_paragraph(normal, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=1.25, before=0, after=0, line_spacing=1.5)

    body = add_or_get_style(doc, "Body Text", WD_STYLE_TYPE.PARAGRAPH)
    set_font(body, 12)
    set_paragraph(body, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=1.25, before=0, after=0, line_spacing=1.5)

    title = styles["Title"]
    set_font(title, 14, bold=True)
    set_paragraph(title, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line_spacing=1.5)

    subtitle = add_or_get_style(doc, "Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    set_font(subtitle, 12)
    set_paragraph(subtitle, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line_spacing=1.5)

    for level, size in [(1, 14), (2, 13), (3, 12), (4, 12)]:
        style = styles[f"Heading {level}"]
        set_font(style, size, bold=True)
        set_paragraph(style, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, before=18 if level == 1 else 12, after=12 if level == 1 else 6, line_spacing=1.5)

    caption = add_or_get_style(doc, "Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_font(caption, 10, bold=False)
    set_paragraph(caption, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, before=6, after=12, line_spacing=1.0)

    block_text = add_or_get_style(doc, "Block Text", WD_STYLE_TYPE.PARAGRAPH)
    set_font(block_text, 11)
    set_paragraph(block_text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=0, left=1.25, before=6, after=6, line_spacing=1.0)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        set_font(style, 12)
        set_paragraph(style, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=0, left=1.25, before=0, after=0, line_spacing=1.5)

    table = styles["Table Grid"]
    table.font.name = "Times New Roman"
    table.font.size = Pt(9)
    try:
        set_cell_margins(table)
    except Exception:
        pass

    doc.add_paragraph("Modelo de estilos do TCC", style="Title")
    doc.add_paragraph("Este arquivo e usado apenas como reference-doc do Pandoc.", style="Body Text")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    main()
