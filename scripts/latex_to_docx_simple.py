import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt


SKIP_COMMANDS = {
    "addcontentsline",
    "captionsetup",
    "clearpage",
    "endfirsthead",
    "endhead",
    "endfoot",
    "endlastfoot",
    "geometry",
    "hypersetup",
    "label",
    "linespread",
    "listoffigures",
    "listoftables",
    "newcommand",
    "noindent",
    "pagebreak",
    "pagenumbering",
    "providecommand",
    "renewcommand",
    "setcounter",
    "setlength",
    "tableofcontents",
    "tightlist",
    "usepackage",
    "vfill",
    "vspace",
    "vspace*",
}

SKIP_ENVIRONMENTS = {
    "center",
    "flushleft",
    "flushright",
    "minipage",
    "titlepage",
}


def strip_comment(line):
    out = []
    escaped = False
    for char in line:
        if char == "%" and not escaped:
            break
        out.append(char)
        escaped = char == "\\" and not escaped
        if char != "\\":
            escaped = False
    return "".join(out)


def extract_braced(text, start):
    if start >= len(text) or text[start] != "{":
        return "", start
    depth = 0
    content = []
    for index in range(start, len(text)):
        char = text[index]
        if char == "{":
            depth += 1
            if depth > 1:
                content.append(char)
        elif char == "}":
            depth -= 1
            if depth == 0:
                return "".join(content), index + 1
            content.append(char)
        else:
            content.append(char)
    return "".join(content), len(text)


def unwrap_commands(text):
    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\\(?:textbf|textit|emph|textsc|texttt|uline|ul|hl)\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\textcolor\{[^{}]*\}\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\(?:Large|LARGE|large|small|footnotesize)\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\url\{([^{}]*)\}", r"\1", text)
    return text


def clean_inline(text):
    text = re.sub(r"\\begin\{(?:center|flushleft|flushright|minipage|titlepage)\}(?:\{[^{}]*\})?", "", text)
    text = re.sub(r"\\end\{(?:center|flushleft|flushright|minipage|titlepage)\}", "", text)
    text = re.sub(r"\\fbox\s*\{", "", text)
    text = text.replace(r"\times", " x ")
    text = text.replace(r"\par", "\n")
    text = unwrap_commands(text)
    text = re.sub(r"\\includegraphics(?:\[[^\]]*\])?\{[^{}]*\}", "", text)
    text = re.sub(r"\\caption\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\label\{[^{}]*\}", "", text)
    text = re.sub(r"\\(?:cite|ref|autoref|pageref)\{([^{}]*)\}", r"[\1]", text)
    text = re.sub(r"\$([^$]*)\$", r"\1", text)
    text = text.replace("``", '"').replace("''", '"')
    replacements = {
        r"\%": "%",
        r"\&": "&",
        r"\_": "_",
        r"\#": "#",
        r"\$": "$",
        r"\{": "{",
        r"\}": "}",
        r"\textless": "<",
        r"\textgreater": ">",
        r"\textquotesingle": "'",
        r"\textordmasculine{}": "o",
        r"~": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("--", "-")
    text = re.sub(r"\\\\", "\n", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^{}]*\})?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" ?\n ?", "\n", text)
    return text.strip()


def command_argument(line, command):
    match = re.search(rf"\\{command}\*?(?:\[[^\]]*\])?\s*", line)
    if not match:
        return None
    pos = match.end()
    while pos < len(line) and line[pos].isspace():
        pos += 1
    if pos < len(line) and line[pos] == "{":
        value, _ = extract_braced(line, pos)
        return clean_inline(value)
    return None


def setup_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"

    return document


def add_paragraph(document, text, style=None, list_style=None):
    text = clean_inline(text)
    if not text:
        return
    paragraphs = [part.strip() for part in text.split("\n") if part.strip()]
    for part in paragraphs:
        paragraph = document.add_paragraph(style=list_style or style)
        paragraph.add_run(part)


def add_image(document, tex_path, image_name):
    image_path = (tex_path.parent / image_name).resolve()
    if not image_path.exists():
        add_paragraph(document, f"[Imagem nao encontrada: {image_name}]")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(5.9))
    except Exception as exc:
        add_paragraph(document, f"[Nao foi possivel inserir imagem {image_name}: {exc}]")


def should_skip_command(line):
    stripped = line.strip()
    for command in SKIP_COMMANDS:
        if re.match(rf"\\{re.escape(command)}(?:\b|\*|\[|\{{)", stripped):
            return True
    if re.match(r"\\(?:begin|end)\{(" + "|".join(re.escape(env) for env in SKIP_ENVIRONMENTS) + r")\}", stripped):
        return True
    return False


def convert(tex_file, output_file):
    tex_path = Path(tex_file).resolve()
    out_path = Path(output_file).resolve()
    document = setup_document()

    raw_text = tex_path.read_text(encoding="utf-8")
    body_match = re.search(r"\\begin\{document\}(.*)\\end\{document\}", raw_text, flags=re.S)
    body = body_match.group(1) if body_match else raw_text

    paragraph_buffer = []

    def flush():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_paragraph(document, " ".join(paragraph_buffer))
            paragraph_buffer = []

    for raw_line in body.splitlines():
        line = strip_comment(raw_line).strip()

        if not line:
            flush()
            continue

        if should_skip_command(line):
            flush()
            continue

        for command, level in [
            ("chapter", 1),
            ("section", 2),
            ("subsection", 3),
            ("subsubsection", 4),
        ]:
            heading = command_argument(line, command)
            if heading:
                flush()
                document.add_heading(heading, level=level)
                line = ""
                break
        if not line:
            continue

        caption = command_argument(line, "caption")
        if caption:
            flush()
            add_paragraph(document, f"Legenda: {caption}", style="Caption")
            continue

        image_match = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^{}]+)\}", line)
        if image_match:
            flush()
            add_image(document, tex_path, image_match.group(1))
            remaining = clean_inline(line[: image_match.start()] + line[image_match.end() :])
            if remaining:
                add_paragraph(document, remaining)
            continue

        item_match = re.match(r"\\item\s*(.*)", line)
        if item_match:
            flush()
            add_paragraph(document, item_match.group(1), list_style="List Bullet")
            continue

        if re.match(r"\\(?:begin|end)\{[^{}]+\}", line):
            flush()
            continue

        cleaned = clean_inline(line)
        if cleaned:
            paragraph_buffer.append(cleaned)

    flush()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Convert a LaTeX file to a simple editable DOCX.")
    parser.add_argument("tex_file")
    parser.add_argument("output_file")
    args = parser.parse_args()
    convert(args.tex_file, args.output_file)


if __name__ == "__main__":
    main()
