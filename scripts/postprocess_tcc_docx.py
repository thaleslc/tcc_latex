import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


RED = RGBColor(255, 0, 0)


def set_run_font(run, size_pt=None, bold=None, color=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size_pt=12, bold=False, italic=False):
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)


def setup_sections(document):
    for section in document.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)


def setup_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for level, size in [(1, 14), (2, 13), (3, 12), (4, 12)]:
        name = f"Heading {level}"
        if name not in styles:
            continue
        style = styles[name]
        set_style_font(style, size, bold=True)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        style.paragraph_format.space_after = Pt(12 if level == 1 else 6)

    if "Caption" in styles:
        caption = styles["Caption"]
        set_style_font(caption, 10)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = Cm(0)
        caption.paragraph_format.line_spacing = 1.0
        caption.paragraph_format.space_before = Pt(6)
        caption.paragraph_format.space_after = Pt(12)


def paragraph_text(paragraph):
    return paragraph.text.strip()


def is_heading(paragraph):
    return paragraph.style and paragraph.style.name.startswith("Heading")


def is_front_matter(index, paragraph):
    text = paragraph_text(paragraph)
    if index < 18:
        return True
    front_titles = {
        "Resumo",
        "Abstract",
        "Agradecimentos",
        "Lista de Figuras",
        "Lista de Tabelas",
        "Sumario",
        "Sumário",
    }
    return text in front_titles


def is_caption_text(text):
    lowered = text.lower()
    return (
        lowered.startswith("figura ")
        or lowered.startswith("tabela ")
        or lowered.startswith("legenda:")
        or "inserir legenda da figura" in lowered
        or "inserir legenda da tabela" in lowered
    )


def add_page_break_before(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:pageBreakBefore")) is None:
        p_pr.append(OxmlElement("w:pageBreakBefore"))


def format_paragraphs(document):
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph_text(paragraph)
        if not text:
            continue

        pf = paragraph.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        if index in {4, 10}:
            add_page_break_before(paragraph)

        if is_heading(paragraph):
            pf.first_line_indent = Cm(0)
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.keep_with_next = True
            if paragraph.style.name == "Heading 1":
                add_page_break_before(paragraph)
                pf.space_before = Pt(0)
                pf.space_after = Pt(18)
        elif is_caption_text(text):
            paragraph.style = document.styles["Caption"] if "Caption" in document.styles else paragraph.style
            pf.first_line_indent = Cm(0)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.line_spacing = 1.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(12)
        elif is_front_matter(index, paragraph):
            pf.first_line_indent = Cm(0)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_after = Pt(12)
            if index == 1:
                pf.space_before = Pt(120)
            elif index == 2:
                pf.space_before = Pt(96)
            elif index == 3:
                pf.space_before = Pt(96)
            elif index in {5, 11}:
                pf.space_before = Pt(48)
            elif index in {7, 13}:
                pf.space_before = Pt(36)
        elif paragraph.style and paragraph.style.name == "Caption":
            pf.first_line_indent = Cm(0)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.line_spacing = 1.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(12)
        elif paragraph.style and paragraph.style.name in {"List Bullet", "List Number"}:
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(1.25)
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            pf.first_line_indent = Cm(1.25)
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in paragraph.runs:
            if run.text:
                set_run_font(run, size_pt=12 if not is_heading(paragraph) else None)
                if is_front_matter(index, paragraph) and index < 18:
                    run.font.color.rgb = RED
                if index in {0, 1, 3, 4, 7, 8, 9, 10, 13, 14}:
                    run.font.size = Pt(14)
                if index in {2, 5, 11}:
                    run.font.size = Pt(16)
                if is_caption_text(text):
                    run.font.size = Pt(10)


def set_cell_margins(table, margin_twips=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)

    for side in ("top", "left", "bottom", "right"):
        node = tbl_cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def format_tables(document):
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_cell_margins(table)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_run_font(run, size_pt=9, bold=True if row_index == 0 else None)


def format_images(document):
    usable_width = Cm(16)  # A4 width minus 3 cm left and 2 cm right margins.
    for shape in document.inline_shapes:
        if shape.width and shape.width > usable_width:
            ratio = usable_width / shape.width
            shape.width = usable_width
            shape.height = int(shape.height * ratio)

    for paragraph in document.paragraphs:
        has_drawing = any(run._element.xpath(".//w:drawing") for run in paragraph.runs)
        if has_drawing:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)


def main():
    if len(sys.argv) != 2:
        raise SystemExit("Uso: python postprocess_tcc_docx.py caminho\\arquivo.docx")

    path = Path(sys.argv[1])
    document = Document(path)
    setup_sections(document)
    setup_styles(document)
    format_paragraphs(document)
    format_tables(document)
    format_images(document)
    document.save(path)


if __name__ == "__main__":
    main()
